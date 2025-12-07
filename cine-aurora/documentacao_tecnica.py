from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import os

def criar_documentacao():
    # Criar documento
    doc = Document()
    
    # Configurar fonte padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Cabeçalho
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "Documentação Técnica - Cine Aurora"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Título
    title = doc.add_heading('Documentação Técnica', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('Plataforma de Streaming', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data
    date = doc.add_paragraph(datetime.now().strftime('%d/%m/%Y'))
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Sumário
    doc.add_heading('Sumário', 1)
    doc.add_paragraph('1. Entrega 01: App (frontend/backend)', style='List Number')
    doc.add_paragraph('2. Entrega 02: Banco de Dados', style='List Number')
    doc.add_paragraph('3. Entrega 03: Armazenamento', style='List Number')
    doc.add_paragraph('4. Entrega 04: Infraestrutura + CI/CD', style='List Number')
    doc.add_paragraph('5. Entrega 05: Monitoramento e Logs', style='List Number')
    doc.add_paragraph('6. Considerações Finais', style='List Number')
    doc.add_page_break()

    # Entrega 01
    doc.add_heading('1. Entrega 01: App (frontend/backend)', level=1)
    doc.add_heading('1.1 Dados da Instância', level=2)
    
    # Tabela de dados da instância
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Detalhe'
    
    data_instancia = [
        ('Provider', 'Google Cloud Platform (Firebase Hosting)'),
        ('Região', 'us-central (padrão Firebase)'),
        ('URL Pública', 'https://cine-aurora-84d97.web.app'),
        ('Tecnologia', 'Firebase Hosting com CDN global'),
        ('Portas Abertas', '80 (HTTP), 443 (HTTPS)'),
        ('Tipo de Hospedagem', 'PaaS (Plataforma como Serviço)')
    ]
    
    for item, desc in data_instancia:
        row_cells = t.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = str(desc)
    
    doc.add_paragraph()
    doc.add_paragraph('Link do App: https://cine-aurora-84d97.web.app', style='Intense Quote')
    doc.add_page_break()

    # Entrega 02
    doc.add_heading('2. Entrega 02: Banco de Dados', level=1)
    doc.add_heading('2.1 Descrição Técnica', level=2)
    
    dados_banco = [
        ('Tipo de Banco', 'Firebase Firestore (NoSQL)'),
        ('Endpoint', 'cine-aurora-84d97.firebaseio.com'),
        ('Região', 'us-central (padrão Firebase)'),
        ('Conexão', 'Firebase Web SDK v9 (modular)'),
        ('Autenticação', 'Firebase Authentication')
    ]
    
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Detalhe'
    
    for item, desc in dados_banco:
        row_cells = t.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = str(desc)
    
    doc.add_heading('2.2 Estrutura do Banco', level=2)
    doc.add_paragraph('Estrutura principal da coleção "series":')
    
    estrutura = '''{
        id: string,                  // ID único da série
        title: string,               // Título da série
        description: string,         // Sinopse
        thumbnailUrl: string,        // URL da miniatura
        bannerUrl: string,           // URL do banner
        releaseYear: number,         // Ano de lançamento
        genres: string[],            // Gêneros
        seasons: [                    // Temporadas
            {
                number: number,      // Número da temporada
                episodes: [          // Episódios
                    {
                        title: string,       // Título do episódio
                        description: string, // Sinopse
                        duration: string,    // Duração (ex: "45m")
                        videoUrl: string,    // URL do vídeo
                        thumbnailUrl: string // Thumbnail do episódio
                    }
                ]
            }
        ]
    }'''
    
    doc.add_paragraph(estrutura, style='Intense Quote')
    doc.add_page_break()

    # Entrega 03
    doc.add_heading('3. Entrega 03: Armazenamento', level=1)
    doc.add_paragraph('Estratégia de armazenamento utilizando links externos:')
    
    doc.add_heading('3.1 Provedores Utilizados', level=2)
    doc.add_paragraph('• Imagens: Google Drive, ImgBB')
    doc.add_paragraph('• Vídeos: Vimeo, YouTube')
    
    doc.add_heading('3.2 Estrutura de Dados', level=2)
    doc.add_paragraph('Exemplo de estrutura com links externos:')
    
    exemplo = '''{
        thumbnailUrl: "https://drive.google.com/...",  // Imagem de capa
        bannerUrl: "https://i.ibb.co/...",           // Banner
        videoUrl: "https://vimeo.com/...",           // Conteúdo de vídeo
    }'''
    
    doc.add_paragraph(exemplo, style='Intense Quote')
    
    doc.add_heading('3.3 Vantagens', level=2)
    doc.add_paragraph('• Custo zero de armazenamento')
    doc.add_paragraph('• Sem preocupação com otimização de mídia')
    doc.add_paragraph('• Escalabilidade ilimitada')
    doc.add_paragraph('• Redução da carga no servidor')
    doc.add_page_break()

    # Entrega 04
    doc.add_heading('4. Entrega 04: Infraestrutura + CI/CD', level=1)
    
    doc.add_heading('4.1 Pipeline CI/CD', level=2)
    doc.add_paragraph('Configuração do GitHub Actions para deploy automático:')
    
    pipeline = '''name: Deploy to Firebase on merge
on:
  push:
    branches: [main]
jobs:
  build_and_deploy:
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v2
      - uses: actions/setup-node@v2
        with:
          node-version: '16'
      - run: npm ci
      - run: npm run build
      - uses: FirebaseExtended/action-hosting-deploy@v0
        with:
          repoToken: '${{ secrets.GITHUB_TOKEN }}'
          firebaseServiceAccount: '${{ secrets.FIREBASE_SERVICE_ACCOUNT }}'
          projectId: cine-aurora-84d97'''
    
    doc.add_paragraph(pipeline, style='Intense Quote')
    
    doc.add_heading('4.2 Fluxo do Pipeline', level=2)
    doc.add_paragraph('1. Trigger: Push para a branch main')
    doc.add_paragraph('2. Setup: Instalação do Node.js e dependências')
    doc.add_paragraph('3. Build: Compilação do projeto React')
    doc.add_paragraph('4. Deploy: Publicação no Firebase Hosting')
    
    doc.add_heading('4.3 Evidências', level=2)
    doc.add_paragraph('• Repositório: https://github.com/SantzL7/CINE-AURORA')
    doc.add_paragraph('• Histórico de implantações disponível no Firebase Console')
    doc.add_page_break()

    # Entrega 05
    doc.add_heading('5. Entrega 05: Monitoramento e Logs', level=1)
    
    doc.add_heading('5.1 Ferramentas Utilizadas', level=2)
    doc.add_paragraph('• Firebase Performance Monitoring')
    doc.add_paragraph('• Firebase Analytics')
    doc.add_paragraph('• Firebase Crashlytics')
    
    doc.add_heading('5.2 Métricas Coletadas', level=2)
    
    doc.add_heading('Desempenho', level=3)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = 'Métrica'
    hdr_cells[1].text = 'Média'
    hdr_cells[2].text = 'Pior 1%'
    
    dados_desempenho = [
        ('Tempo de Inicialização (app_start)', '1.1s', '3.7s'),
        ('Tempo de Renderização (screen_view)', '1.3s', '4.5s')
    ]
    
    for metrica, media, pior in dados_desempenho:
        row_cells = t.add_row().cells
        row_cells[0].text = metrica
        row_cells[1].text = media
        row_cells[2].text = pior
    
    doc.add_paragraph()
    doc.add_heading('Uso', level=3)
    doc.add_paragraph('• Usuários ativos diários')
    doc.add_paragraph('• Eventos de navegação')
    doc.add_paragraph('• Taxa de retenção')
    
    doc.add_heading('Estabilidade', level=3)
    doc.add_paragraph('• Taxa de falhas')
    doc.add_paragraph('• Erros não tratados')
    doc.add_paragraph('• Dispositivos afetados')
    
    doc.add_heading('5.3 Acesso', level=2)
    doc.add_paragraph('Console do Firebase: https://console.firebase.google.com/')
    doc.add_paragraph('Projeto: Cine Aurora')
    
    # Considerações Finais
    doc.add_heading('6. Considerações Finais', level=1)
    
    doc.add_heading('Próximos Passos', level=2)
    doc.add_paragraph('1. Implementar testes automatizados')
    doc.add_paragraph('2. Configurar alertas de monitoramento')
    doc.add_paragraph('3. Otimizar desempenho das telas mais lentas')
    
    doc.add_heading('Contato', level=2)
    doc.add_paragraph('• Repositório: https://github.com/SantzL7/CINE-AURORA')
    doc.add_paragraph('• Aplicação: https://cine-aurora-84d97.web.app')
    
    # Salvar documento
    output_path = os.path.join(os.path.expanduser('~'), 'Desktop', 'Documentacao_CineAurora.docx')
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    try:
        output = criar_documentacao()
        print(f"Documentação gerada com sucesso em: {output}")
    except Exception as e:
        print(f"Erro ao gerar documentação: {str(e)}")
